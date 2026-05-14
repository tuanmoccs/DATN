import base64
import io
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from langchain_core.messages import HumanMessage

from app.core.dependencies import get_llm

PDF_TEXT_LAYER_MIN_CHARS = 30
PDF_OCR_MAX_PAGES = 30
PDF_TEXT_MAX_SUSPICIOUS_RATIO = 0.08
PDF_TEXT_MAX_CYRILLIC_RATIO = 0.02
PDF_TEXT_MAX_CONTROL_RATIO = 0.01


def extract_text_from_file(file_path: str, content_type: str) -> str:
    """Extract text content from uploaded files."""
    path = Path(file_path)

    if content_type == "pdf":
        return _extract_pdf(path)
    elif content_type == "docx":
        return _extract_docx(path)
    elif content_type == "txt":
        return _extract_txt(path)
    else:
        raise ValueError(f"Unsupported content type: {content_type}")


def extract_text_from_bytes(file_bytes: bytes, content_type: str) -> str:
    """Extract text from file bytes."""
    if content_type == "pdf":
        return _extract_pdf_bytes(file_bytes)
    elif content_type == "docx":
        return _extract_docx_bytes(file_bytes)
    elif content_type == "txt":
        return _decode_txt_bytes(file_bytes)
    else:
        raise ValueError(f"Unsupported content type: {content_type}")


async def extract_text_from_bytes_with_fallback(file_bytes: bytes, content_type: str) -> str:
    """Extract text and use OCR fallback for PDF pages without a readable text layer."""
    if content_type != "pdf":
        return extract_text_from_bytes(file_bytes, content_type)

    return await _extract_pdf_bytes_with_fallback(file_bytes)


def _extract_pdf(path: Path) -> str:
    doc = fitz.open(str(path))
    text_parts = []
    for page in doc:
        text_parts.append(page.get_text())
    doc.close()
    return "\n".join(text_parts)


def _extract_pdf_bytes(file_bytes: bytes) -> str:
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    text_parts = []
    for page in doc:
        text_parts.append(page.get_text())
    doc.close()
    return "\n".join(text_parts)


async def _extract_pdf_bytes_with_fallback(file_bytes: bytes) -> str:
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    text_parts: list[str] = []

    try:
        for page_index, page in enumerate(doc):
            page_text = page.get_text().strip()

            if _is_readable_pdf_text(page_text):
                text_parts.append(page_text)
                continue

            if page_index >= PDF_OCR_MAX_PAGES:
                if page_text:
                    text_parts.append(page_text)
                continue

            try:
                ocr_text = await _ocr_pdf_page(page)
            except Exception:
                ocr_text = ""
            if ocr_text.strip():
                text_parts.append(ocr_text.strip())
            elif page_text:
                text_parts.append(page_text)
    finally:
        doc.close()

    return "\n\n".join(part for part in text_parts if part.strip())


def _is_readable_pdf_text(text: str) -> bool:
    normalized = text.strip()
    if len(normalized) < PDF_TEXT_LAYER_MIN_CHARS:
        return False

    total = len(normalized)
    if total == 0:
        return False

    control_chars = sum(1 for char in normalized if _is_control_char(char))
    cyrillic_chars = sum(1 for char in normalized if "\u0400" <= char <= "\u04ff")
    suspicious_chars = sum(1 for char in normalized if _is_suspicious_pdf_char(char))

    return (
        control_chars / total <= PDF_TEXT_MAX_CONTROL_RATIO
        and cyrillic_chars / total <= PDF_TEXT_MAX_CYRILLIC_RATIO
        and suspicious_chars / total <= PDF_TEXT_MAX_SUSPICIOUS_RATIO
    )


def _is_control_char(char: str) -> bool:
    return ord(char) < 32 and char not in "\n\r\t"


def _is_suspicious_pdf_char(char: str) -> bool:
    codepoint = ord(char)
    if _is_control_char(char):
        return True
    if "\u0400" <= char <= "\u04ff":
        return True
    if char in {"�", "□"}:
        return True
    return codepoint in {0xFFFD, 0x25A1}


def _extract_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])


def _extract_docx_bytes(file_bytes: bytes) -> str:
    doc = DocxDocument(io.BytesIO(file_bytes))
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])


def _extract_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def _decode_txt_bytes(file_bytes: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "utf-16"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue

    raise UnicodeDecodeError("text", file_bytes, 0, 1, "Unsupported text encoding")


async def _ocr_pdf_page(page: fitz.Page) -> str:
    llm = get_llm()

    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
    image_b64 = base64.b64encode(pix.tobytes("png")).decode("ascii")

    message = HumanMessage(content=[
        {
            "type": "text",
            "text": (
                "Extract all readable text from this Vietnamese textbook page. "
                "Return plain text only. Preserve headings, formulas, bullet points, "
                "page numbers, questions, and Vietnamese diacritics. Do not summarize. "
                "Ignore decorative images unless they contain labels or captions."
            ),
        },
        {
            "type": "image_url",
            "image_url": {
                "url": f"data:image/png;base64,{image_b64}",
            },
        },
    ])

    response = await llm.ainvoke([message])

    if isinstance(response.content, str):
        return response.content.strip()

    if isinstance(response.content, list):
        text_parts = []
        for item in response.content:
            if isinstance(item, dict) and item.get("type") == "text":
                text_parts.append(item.get("text", ""))
        return "\n".join(part.strip() for part in text_parts if part.strip())

    return ""
